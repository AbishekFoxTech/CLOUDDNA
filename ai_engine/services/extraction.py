import io

from django.conf import settings

try:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
except ImportError:
    pytesseract = None

OCR_TARGET_DIMENSION = 1600  # upscale smaller images up to roughly this size


class TextExtractionError(Exception):
    """
    Raised when a document's text genuinely cannot be extracted: a corrupt
    file, an unsupported/missing OCR engine, or unreadable content. Callers
    (the pipeline) catch this and mark the document as AI Processing Failed
    with a user-friendly message, rather than letting it crash the request.
    """


def extract_text(file_type, file_bytes):
    file_type = (file_type or '').lower().lstrip('.')

    if file_type == 'pdf':
        return _extract_pdf(file_bytes)
    if file_type == 'docx':
        return _extract_docx(file_bytes)
    if file_type == 'txt':
        return _extract_txt(file_bytes)
    if file_type in {'png', 'jpg', 'jpeg'}:
        return _extract_image(file_bytes)

    raise TextExtractionError(f'No text extraction strategy available for ".{file_type}" files.')


def _extract_txt(file_bytes):
    try:
        return file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        return file_bytes.decode('latin-1', errors='ignore')


def _extract_docx(file_bytes):
    try:
        import docx
    except ImportError as exc:
        raise TextExtractionError('python-docx is not installed.') from exc

    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise TextExtractionError(f'Could not read this DOCX file: {exc}') from exc

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return '\n'.join(paragraphs)


def _extract_pdf(file_bytes):
    try:
        import pdfplumber
    except ImportError as exc:
        raise TextExtractionError('pdfplumber is not installed.') from exc

    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ''
                if page_text.strip():
                    text_parts.append(page_text)
    except Exception as exc:
        raise TextExtractionError(f'Could not read this PDF file: {exc}') from exc

    text = '\n'.join(text_parts).strip()
    if text:
        return text

    # No embedded text found - likely a scanned/image-only PDF. Fall back to
    # rendering each page as an image and running OCR on it.
    return _extract_pdf_via_ocr(file_bytes)


def _extract_pdf_via_ocr(file_bytes):
    try:
        import fitz  # PyMuPDF
        from PIL import Image
    except ImportError as exc:
        raise TextExtractionError('OCR fallback dependencies are not installed.') from exc

    text_parts = []
    try:
        pdf_doc = fitz.open(stream=file_bytes, filetype='pdf')
        max_pages = min(len(pdf_doc), settings.AI_OCR_MAX_PDF_PAGES)
        for page_index in range(max_pages):
            page = pdf_doc[page_index]
            pixmap = page.get_pixmap(dpi=200)
            image = Image.frombytes('RGB', [pixmap.width, pixmap.height], pixmap.samples)
            page_text = _best_effort_ocr(image, thorough=False)
            if page_text.strip():
                text_parts.append(page_text)
        pdf_doc.close()
    except _TesseractNotFound as exc:
        raise TextExtractionError(
            'This looks like a scanned PDF, but the OCR engine (Tesseract) '
            'is not installed on the server.'
        ) from exc
    except Exception as exc:
        raise TextExtractionError(f'Could not OCR this scanned PDF: {exc}') from exc

    return '\n'.join(text_parts).strip()


def _extract_image(file_bytes):
    try:
        from PIL import Image
    except ImportError as exc:
        raise TextExtractionError('pytesseract/Pillow are not installed.') from exc

    try:
        image = Image.open(io.BytesIO(file_bytes))
        image.load()
        if image.mode not in ('RGB', 'L'):
            image = image.convert('RGB')
    except Exception as exc:
        raise TextExtractionError(f'Could not read this image file: {exc}') from exc

    try:
        return _best_effort_ocr(image, thorough=True)
    except _TesseractNotFound as exc:
        raise TextExtractionError(
            'The OCR engine (Tesseract) is not installed on the server, so '
            'text could not be extracted from this image.'
        ) from exc
    except Exception as exc:
        raise TextExtractionError(f'OCR failed for this image: {exc}') from exc


class _TesseractNotFound(Exception):
    """Internal sentinel so callers don't need to import pytesseract directly."""


def _otsu_threshold(grayscale_image):
    """
    Binarizes a grayscale image using Otsu's method: picks the threshold
    that best separates two intensity populations (dark text vs. a lighter
    background/pattern) by maximizing between-class variance, rather than
    guessing a fixed cutoff. Implemented directly since no OpenCV/
    scikit-image dependency is in this project.
    """
    import numpy as np

    histogram = np.array(grayscale_image.histogram(), dtype=float)
    total = histogram.sum()
    sum_total = np.dot(np.arange(256), histogram)

    weight_background = 0.0
    sum_background = 0.0
    best_variance = 0.0
    threshold = 127

    for level in range(256):
        weight_background += histogram[level]
        if weight_background == 0:
            continue
        weight_foreground = total - weight_background
        if weight_foreground == 0:
            break

        sum_background += level * histogram[level]
        mean_background = sum_background / weight_background
        mean_foreground = (sum_total - sum_background) / weight_foreground

        between_class_variance = (
            weight_background * weight_foreground * (mean_background - mean_foreground) ** 2
        )
        if between_class_variance > best_variance:
            best_variance = between_class_variance
            threshold = level

    return grayscale_image.point(lambda p: 255 if p > threshold else 0)


def _ocr_candidates(image):
    """
    Yields (PIL.Image, tesseract config) pairs to try. Real-world photos of
    printed material (glare, background patterns, uneven lighting - e.g. a
    photographed business card) often OCR poorly with Tesseract's defaults,
    so we try a small, cheap set of preprocessing variants and let
    confidence scoring (see _best_effort_ocr) pick the winner rather than
    guessing a single fixed pipeline.
    """
    from PIL import Image, ImageOps

    yield image, ''

    grayscale = ImageOps.autocontrast(image.convert('L'))
    yield grayscale, ''

    largest_dimension = max(grayscale.size)
    if largest_dimension < OCR_TARGET_DIMENSION:
        scale = OCR_TARGET_DIMENSION / largest_dimension
        upscaled = grayscale.resize(
            (int(grayscale.width * scale), int(grayscale.height * scale)), Image.LANCZOS,
        )
    else:
        upscaled = grayscale
    yield upscaled, '--psm 6'

    try:
        binarized = _otsu_threshold(upscaled)
        yield binarized, '--psm 6'
    except Exception:
        pass


def _reconstruct_text(ocr_data):
    """Groups Tesseract's word-level output back into lines, preserving
    layout better than a flat join of every recognized word."""
    lines = {}
    for i, word in enumerate(ocr_data['text']):
        if not word.strip():
            continue
        key = (ocr_data['block_num'][i], ocr_data['par_num'][i], ocr_data['line_num'][i])
        lines.setdefault(key, []).append(word)
    return '\n'.join(' '.join(words) for words in lines.values())


def _best_effort_ocr(image, thorough=True):
    """
    Runs OCR against a small set of preprocessing candidates and keeps
    whichever produced the highest average word-confidence. `thorough=False`
    limits this to two candidates (used for scanned-PDF pages, where a
    single document may have many pages).
    """
    if pytesseract is None:
        raise TextExtractionError('pytesseract is not installed.')
    from pytesseract import Output

    best_text = ''
    best_score = -1.0
    tried_any = False

    for index, (candidate_image, config) in enumerate(_ocr_candidates(image)):
        if not thorough and index >= 2:
            break
        try:
            data = pytesseract.image_to_data(candidate_image, config=config, output_type=Output.DICT)
            tried_any = True
        except pytesseract.TesseractNotFoundError as exc:
            raise _TesseractNotFound() from exc
        except Exception:
            continue

        confidences = [float(c) for c in data.get('conf', []) if c not in ('-1', -1)]
        if not confidences:
            continue

        text = _reconstruct_text(data)
        if not text.strip():
            continue

        avg_confidence = sum(confidences) / len(confidences)
        if avg_confidence > best_score:
            best_score = avg_confidence
            best_text = text

    if not tried_any:
        raise TextExtractionError('OCR engine did not return a usable result for this file.')

    return best_text.strip()
